"""
JobMatch — AI-powered job application assistant
Built for Kahini Gandhi (Senior PM/BA, PSM I, 10+ yrs)

Tabs:
  1. Profile     - store/edit resume text
  2. Job Feed    - pull live listings (Adzuna API), score fit, tailor, apply, log
  3. Analyze Fit - one-off: paste any job posting -> fit score + gap analysis
  4. Tailor      - generate a JD-tailored resume -> export as .docx
  5. Cover       - generate a matching cover letter -> export as .docx
  6. Dashboard   - every application logged, status, and the exact tailored
                   resume/cover letter used for it
  7. Email Check - placeholder (backlog item, not yet implemented)
"""

import os
import io
import json
import sqlite3
import datetime as dt

import streamlit as st
import pandas as pd
import requests
from bs4 import BeautifulSoup

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import anthropic

# ----------------------------------------------------------------------------
# Config
# ----------------------------------------------------------------------------

st.set_page_config(page_title="JobMatch", page_icon="🎯", layout="wide")

DEFAULT_MODEL = "claude-sonnet-4-6"
DB_PATH = "jobmatch.db"

CANDIDATE_NAME = "Kahini Gandhi"
CANDIDATE_CONTACT = "gandhi.kahini91@gmail.com | (484) 798-1696 | Ontario, Canada"

DEFAULT_RESUME = """KAHINI GANDHI
gandhi.kahini91@gmail.com | (484) 798-1696 | Ontario, Canada | linkedin.com/in/kahini-gandhi

Project Manager / Lead Business Analyst | Salesforce - ServiceNow - Agile - ERP
10+ Years - PSM I Certified - Healthcare - Government - Technology

PROFESSIONAL SUMMARY
Senior PM/BA with 10+ years delivering enterprise IT solutions across Salesforce, ServiceNow, ERP, and healthcare domains. Proven record managing $1M+ platform projects, working with 40-50 person cross-functional teams, and converting ambiguous business needs into developer-ready requirements. Expert in full SDLC governance, Agile delivery, and CRM/ITSM implementation from discovery through UAT.

CORE COMPETENCIES
Project Management (PM), Salesforce CRM / Sales Cloud, ServiceNow ITSM, Lead Business Analysis, Agile / Scrum / SAFe / Kanban, Product Ownership & Roadmap, Requirements Engineering, Stakeholder & Vendor Management, ERP & GAP Analysis, SDLC / UAT Governance, Risk & Change Management, BI & Data Analysis (Tableau), Process Automation, HIPAA / SOX Compliance, IT Governance

PROFESSIONAL EXPERIENCE

Lead Business Analyst / Process Analyst | Spruce Technology -- Clifton, NJ | May 2022 -- Present
Salesforce - ServiceNow - JIRA - Tableau - SharePoint - ERP
- Accelerated complaint resolution: Re-engineered Salesforce CRM workflows, eliminating manual handoff steps and automating escalation triggers, cutting resolution time and improving customer satisfaction.
- Delivered ERP migration on schedule: Led GAP analysis across multiple business units; produced item prioritized remediation backlog that kept implementation on track with zero critical post-go-live defects.
- Reduced ServiceNow incident backlog: Defined requirements for ITSM escalation workflows, authored acceptance criteria for configuration stories, and signed off UAT.
- Delivered executive KPI reporting: Designed and deployed a multi-source dashboard (Azure DevOps / Tableau, structured for Power BI consumption) tracking complaint handling, risk exposure, and team performance across programs, replacing manual weekly reports and enabling real-time visibility for senior leadership.
- Influenced senior stakeholders across program lifecycle: Served as primary liaison across business, compliance, and IT; facilitated cross-functional workshops each sprint and presented program status, risk exposure, and recommendations directly to VP and Director-level stakeholders.
- Leveraged Salesforce Einstein AI: Explored and documented use cases for Einstein AI-powered case classification and next-best-action recommendations within Salesforce Service Cloud; contributed to an internal feasibility assessment for AI-assisted workflow automation.
- Embedded IT risk controls into delivery: Maintained risk registers and authored risk-based acceptance criteria across projects involving regulated data (HIPAA, SOX); flagged compliance and IT risk gaps to VP-level leadership and tracked remediation to closure.

Product Owner | iSolufy LLC -- Zephyrhills, FL | Sep 2020 -- Jan 2022
Salesforce - JIRA - Confluence - Balsamiq - Axure RP - HIPAA PHI Systems
- Eliminated data transfer error rate: Authored user stories and defined data validation rules for a Salesforce-integrated platform, eliminating manual reconciliation effort.
- Managed product backlogs: Owned full backlog lifecycle for Agile team across multiple release cycles; maintained sprint commitment accuracy.
- Achieved HIPAA compliance, zero audit findings: Embedded privacy controls into acceptance criteria for all PHI-touching features; coordinated internal audits with no findings and zero regulatory violations.
- Expanded accessible user base: Partnered with UX on WCAG testing, identifying and resolving critical accessibility gaps pre-launch.

Business Systems Analyst | Pharmfood Group -- Tampa, FL | Jul 2016 -- Jun 2020
Azure DevOps - ServiceNow - JIRA - MS Visio - SQL - MS Project
- Reduced process cycle time: Analyzed multiple departments, identified workflow bottlenecks, and redesigned ServiceNow workflows.
- Built ADO governance framework: Designed custom Azure DevOps board architecture with workflow states and impediment tracking, adopted by multiple teams.
- Delivered major releases with minimum defect escape: Coordinated full UAT lifecycle -- test plans, defect triage, and stakeholder sign-off -- achieving on-time deployment for major releases.
- Zero spec-related defects across requirements: Translated business requirements into technical specifications.

EDUCATION
M.S., Computer Science -- Texas A&M University, Kingsville, TX (2014-2016)
B.E., Computer Engineering -- Gujarat Technological University, India (2009-2013)

CERTIFICATIONS
Professional Scrum Master (PSM I) -- Scrum.org
Introduction to Prompt Engineering for Generative AI -- LinkedIn Learning

TECHNICAL SKILLS
CRM / ITSM: Salesforce (Sales Cloud, Service Cloud, Einstein AI, Agentforce, Workflows, Reports), ServiceNow (ITSM, Incident, Problem, Change, CSM)
PM / ALM Tools: JIRA, Azure DevOps, MS Project, Confluence, SharePoint, Trello, Power BI
Methodologies: Agile, Scrum, SAFe, Kanban, Waterfall, Hybrid, SDLC, PDLC
BA Techniques: User Stories, Use Cases, BRD, FRD, Process Mapping, Wireframing, GAP / Fit-Gap Analysis, Impact Analysis
BI & Analytics: Tableau, Amazon QuickSight, Google Analytics, KPI Dashboards
Data / Integration: SQL, XML, JSON, REST API, HTML, CSS, JavaScript
Design / UX: Balsamiq, Axure RP, Adobe XD, Wireframes, Prototyping
Compliance: HIPAA, SOX, Regulatory Reporting, Risk & Change Management

INDUSTRY EXPERIENCE
Healthcare, Pharmaceutical, Government / Housing Authority, Technology / SaaS, Automobile/Transit, Financial Services
"""

# ----------------------------------------------------------------------------
# Secrets / API helpers
# ----------------------------------------------------------------------------

def get_secret(key):
    val = None
    try:
        val = st.secrets.get(key)
    except Exception:
        val = None
    if not val:
        val = os.environ.get(key)
    return val


@st.cache_resource
def get_client():
    api_key = get_secret("ANTHROPIC_API_KEY")
    if not api_key:
        st.error(
            "No Anthropic API key found. Add `ANTHROPIC_API_KEY` to this app's "
            "Secrets (Streamlit Community Cloud: Settings -> Secrets)."
        )
        st.stop()
    return anthropic.Anthropic(api_key=api_key)


FIT_SCHEMA = {
    "name": "fit_analysis",
    "description": "Analyze fit between a resume and a job description.",
    "input_schema": {
        "type": "object",
        "properties": {
            "fit_score": {"type": "integer", "description": "0-100 overall fit score"},
            "role_title": {"type": "string"},
            "company_name": {"type": "string", "description": "Best guess of company name, or 'Unknown'"},
            "matched_skills": {"type": "array", "items": {"type": "string"}},
            "missing_skills": {"type": "array", "items": {"type": "string"}},
            "summary": {"type": "string", "description": "2-3 sentence honest assessment of fit"},
            "recommendation": {"type": "string", "enum": ["Strong match - apply", "Possible match - review gaps", "Weak match - likely skip"]},
        },
        "required": ["fit_score", "role_title", "company_name", "matched_skills", "missing_skills", "summary", "recommendation"],
    },
}

TAILOR_SCHEMA = {
    "name": "tailored_resume",
    "description": "Produce a tailored resume summary and bullet rewrites based on a job description. Never invent employers, titles, dates, or skills not present in the original resume.",
    "input_schema": {
        "type": "object",
        "properties": {
            "tailored_headline": {"type": "string"},
            "tailored_summary": {"type": "string"},
            "keyword_notes": {"type": "array", "items": {"type": "string"}},
            "experience_rewrites": {
                "type": "array",
                "items": {
                    "type": "object",
                    "properties": {
                        "role_company": {"type": "string"},
                        "bullets": {"type": "array", "items": {"type": "string"}},
                    },
                    "required": ["role_company", "bullets"],
                },
            },
        },
        "required": ["tailored_headline", "tailored_summary", "keyword_notes", "experience_rewrites"],
    },
}

COVER_LETTER_SCHEMA = {
    "name": "cover_letter",
    "description": "Write a concise, specific cover letter grounded only in the candidate's real resume content.",
    "input_schema": {
        "type": "object",
        "properties": {
            "greeting": {"type": "string"},
            "body_paragraphs": {"type": "array", "items": {"type": "string"}, "minItems": 3, "maxItems": 4},
            "closing": {"type": "string"},
        },
        "required": ["greeting", "body_paragraphs", "closing"],
    },
}


def call_tool(system, user_prompt, schema, model=None, max_tokens=3000):
    client = get_client()
    try:
        resp = client.messages.create(
            model=model or DEFAULT_MODEL,
            max_tokens=max_tokens,
            system=system,
            tools=[schema],
            tool_choice={"type": "tool", "name": schema["name"]},
            messages=[{"role": "user", "content": user_prompt}],
        )
        for block in resp.content:
            if block.type == "tool_use":
                return block.input
        return None
    except Exception as e:
        st.error(f"AI request failed: {e}")
        return None


def fetch_job_posting(url):
    try:
        r = requests.get(url, timeout=15, headers={"User-Agent": "Mozilla/5.0"})
        r.raise_for_status()
        soup = BeautifulSoup(r.text, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = soup.get_text(separator="\n")
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        return "\n".join(lines)[:8000]
    except Exception as e:
        st.error(f"Couldn't fetch that URL: {e}")
        return None


def search_jooble(keywords, location, page=1):
    api_key = get_secret("JOOBLE_API_KEY")
    if not api_key:
        st.error(
            "No Jooble API key found. Add `JOOBLE_API_KEY` to this app's Secrets. "
            "Free account: jooble.org/api/about"
        )
        return []
    url = f"https://jooble.org/api/{api_key}"
    try:
        r = requests.post(url, json={"keywords": keywords, "location": location, "page": str(page)}, timeout=20)
        r.raise_for_status()
        return r.json().get("jobs", [])
    except Exception as e:
        st.error(f"Jooble search failed: {e}")
        return []


def fetch_greenhouse_board(board_token):
    try:
        r = requests.get(f"https://boards-api.greenhouse.io/v1/boards/{board_token}/jobs", timeout=15)
        r.raise_for_status()
        return r.json().get("jobs", [])
    except Exception as e:
        st.error(f"Greenhouse fetch failed for '{board_token}': {e}")
        return []


def fetch_lever_board(company_slug):
    try:
        r = requests.get(f"https://api.lever.co/v0/postings/{company_slug}?mode=json", timeout=15)
        r.raise_for_status()
        return r.json()
    except Exception as e:
        st.error(f"Lever fetch failed for '{company_slug}': {e}")
        return []


QUICK_LINKS = [
    ("LinkedIn Jobs", "linkedin", "https://www.linkedin.com/jobs/search/?keywords={kw}&location={loc}"),
    ("Indeed", "indeed", "https://ca.indeed.com/jobs?q={kw}&l={loc}"),
    ("Job Bank Canada (govt + general)", "jobbank", "https://www.jobbank.gc.ca/jobsearch/jobsearch?searchstring={kw}&locationstring={loc}"),
    ("Government of Canada Jobs (public service)", "gcjobs", "https://emploisfp-psjobs.cfp-psc.gc.ca/psrs-srfp/applicant/page1800"),
    ("Deloitte Canada Careers", "deloitte", "https://www2.deloitte.com/ca/en/careers/careers.html"),
    ("PwC Canada Careers", "pwc", "https://www.pwc.com/ca/en/careers.html"),
    ("EY Canada Careers", "ey", "https://careers.ey.com/ey/"),
    ("KPMG Canada Careers", "kpmg", "https://kpmg.com/ca/en/home/careers.html"),
    ("Robert Half Canada", "roberthalf", "https://www.roberthalf.com/ca/en/find-jobs?q={kw}"),
    ("CGI Careers", "cgi", "https://www.cgi.com/en/careers"),
    ("TD Bank Careers", "td", "https://careers.td.com/"),
    ("CIBC Careers", "cibc", "https://cibc.wd3.myworkdayjobs.com/search?q={kw}"),
    ("TCS Careers", "tcs", "https://www.tcs.com/careers"),
    ("Google Careers", "google", "https://www.google.com/about/careers/applications/jobs/results/?q={kw}"),
    ("Amazon Jobs", "amazon", "https://www.amazon.jobs/en/search?base_query={kw}"),
    ("Meta Careers", "meta", "https://www.metacareers.com/jobs/?q={kw}"),
    ("Apple Jobs", "apple", "https://jobs.apple.com/en-us/search?search={kw}"),
    ("Netflix Jobs", "netflix", "https://explore.jobs.netflix.net/careers?query={kw}"),
]


def search_adzuna(keywords, location, country="ca", page=1, results_per_page=20):
    app_id = get_secret("ADZUNA_APP_ID")
    app_key = get_secret("ADZUNA_APP_KEY")
    if not app_id or not app_key:
        st.error(
            "No Adzuna API credentials found. Add `ADZUNA_APP_ID` and `ADZUNA_APP_KEY` "
            "to this app's Secrets. Free account: developer.adzuna.com"
        )
        return []
    url = f"https://api.adzuna.com/v1/api/jobs/{country}/search/{page}"
    params = {
        "app_id": app_id,
        "app_key": app_key,
        "results_per_page": results_per_page,
        "what": keywords,
        "where": location,
        "content-type": "application/json",
    }
    try:
        r = requests.get(url, params=params, timeout=20)
        r.raise_for_status()
        return r.json().get("results", [])
    except Exception as e:
        st.error(f"Adzuna search failed: {e}")
        return []


# ----------------------------------------------------------------------------
# Persistence (SQLite)
# ----------------------------------------------------------------------------

@st.cache_resource
def get_db():
    conn = sqlite3.connect(DB_PATH, check_same_thread=False)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS applications (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company TEXT, role TEXT, url TEXT, fit_score INTEGER, status TEXT,
            date_applied TEXT, date_updated TEXT, notes TEXT,
            jd_text TEXT, tailor_json TEXT, cover_json TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS job_feed (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT, company TEXT, location TEXT, salary TEXT, url TEXT UNIQUE,
            description TEXT, source TEXT, fit_score INTEGER, fit_json TEXT,
            tailor_json TEXT, cover_json TEXT, status TEXT, fetched_at TEXT
        )
    """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS company_watchlist (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            company_name TEXT, platform TEXT, board_token TEXT
        )
    """)
    conn.commit()
    return conn


def get_watchlist_df(conn):
    return pd.read_sql("SELECT * FROM company_watchlist ORDER BY id DESC", conn)


def add_watchlist_company(conn, company_name, platform, board_token):
    conn.execute(
        "INSERT INTO company_watchlist (company_name, platform, board_token) VALUES (?, ?, ?)",
        (company_name, platform, board_token),
    )
    conn.commit()


def remove_watchlist_company(conn, row_id):
    conn.execute("DELETE FROM company_watchlist WHERE id = ?", (row_id,))
    conn.commit()


def upsert_job_feed_row(conn, job):
    existing = conn.execute("SELECT id FROM job_feed WHERE url = ?", (job["url"],)).fetchone()
    if existing:
        return existing[0]
    cur = conn.execute(
        "INSERT INTO job_feed (title, company, location, salary, url, description, source, status, fetched_at) "
        "VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)",
        (job["title"], job["company"], job["location"], job.get("salary", ""), job["url"],
         job["description"], job.get("source", "Adzuna"), "new", dt.datetime.now().isoformat()),
    )
    conn.commit()
    return cur.lastrowid


def update_job_feed_row(conn, job_id, **fields):
    sets = ", ".join(f"{k} = ?" for k in fields)
    conn.execute(f"UPDATE job_feed SET {sets} WHERE id = ?", (*fields.values(), job_id))
    conn.commit()


def get_job_feed_df(conn):
    return pd.read_sql("SELECT * FROM job_feed ORDER BY id DESC", conn)


def insert_application(conn, **fields):
    cols = ", ".join(fields.keys())
    placeholders = ", ".join("?" for _ in fields)
    cur = conn.execute(f"INSERT INTO applications ({cols}) VALUES ({placeholders})", tuple(fields.values()))
    conn.commit()
    return cur.lastrowid


def update_application(conn, app_id, **fields):
    sets = ", ".join(f"{k} = ?" for k in fields)
    conn.execute(f"UPDATE applications SET {sets} WHERE id = ?", (*fields.values(), app_id))
    conn.commit()


def get_applications_df(conn):
    return pd.read_sql("SELECT * FROM applications ORDER BY id DESC", conn)


# ----------------------------------------------------------------------------
# Docx export helpers
# ----------------------------------------------------------------------------

def export_tailored_resume_docx(tailor_data):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    p = doc.add_paragraph()
    run = p.add_run(CANDIDATE_NAME)
    run.bold = True
    run.font.size = Pt(16)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    headline = doc.add_paragraph()
    hr = headline.add_run(tailor_data.get("tailored_headline", ""))
    hr.italic = True
    hr.font.size = Pt(11)
    headline.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    doc.add_paragraph().add_run("PROFESSIONAL SUMMARY").bold = True
    doc.add_paragraph(tailor_data.get("tailored_summary", ""))

    doc.add_paragraph()
    doc.add_paragraph().add_run("PROFESSIONAL EXPERIENCE").bold = True

    for entry in tailor_data.get("experience_rewrites", []):
        rp = doc.add_paragraph()
        rp.add_run(entry.get("role_company", "")).bold = True
        for bullet in entry.get("bullets", []):
            doc.add_paragraph(bullet, style="List Bullet")
        doc.add_paragraph()

    note = doc.add_paragraph()
    note.add_run(
        "Education, certifications, and technical skills carried over unchanged from original resume."
    ).italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


def export_cover_letter_docx(cover_data):
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    p = doc.add_paragraph()
    run = p.add_run(CANDIDATE_NAME)
    run.bold = True
    run.font.size = Pt(14)

    doc.add_paragraph(CANDIDATE_CONTACT)
    doc.add_paragraph(dt.date.today().strftime("%B %d, %Y"))
    doc.add_paragraph()

    doc.add_paragraph(cover_data.get("greeting", "Dear Hiring Manager,"))
    for para in cover_data.get("body_paragraphs", []):
        doc.add_paragraph(para)
    doc.add_paragraph(cover_data.get("closing", "Sincerely,"))
    doc.add_paragraph(CANDIDATE_NAME)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf


# ----------------------------------------------------------------------------
# Session state
# ----------------------------------------------------------------------------

conn = get_db()

if "resume_text" not in st.session_state:
    st.session_state.resume_text = DEFAULT_RESUME
if "jd_text" not in st.session_state:
    st.session_state.jd_text = ""
if "fit_result" not in st.session_state:
    st.session_state.fit_result = None
if "tailor_result" not in st.session_state:
    st.session_state.tailor_result = None
if "cover_result" not in st.session_state:
    st.session_state.cover_result = None

# ----------------------------------------------------------------------------
# UI
# ----------------------------------------------------------------------------

st.title("🎯 JobMatch")
st.caption("Job application assistant for Kahini Gandhi — find openings, check fit, tailor resumes, track every application.")

with st.sidebar:
    st.header("Navigation")
    page = st.radio(
        "Go to",
        ["1. Profile", "2. Job Feed", "3. Analyze Fit", "4. Tailor Resume", "5. Cover Letter", "6. Dashboard", "7. Email Check"],
        label_visibility="collapsed",
    )
    st.divider()
    st.caption("Nothing is ever submitted automatically. Kahini reviews and clicks Apply herself on the real job site.")

# --- Page 1: Profile ---
if page == "1. Profile":
    st.subheader("Resume Profile")
    st.write("Source of truth for every tailored resume and cover letter. Edit it anytime her real resume changes.")
    st.session_state.resume_text = st.text_area("Resume text", value=st.session_state.resume_text, height=500)
    st.success("Saved automatically as you edit — this feeds every other tab.")

# --- Page 2: Job Feed ---
elif page == "2. Job Feed":
    st.subheader("Job Feed")
    st.caption("Pulls live listings from Adzuna + Jooble, plus any watchlisted companies. Review, check fit, tailor, then click through to apply on the real posting.")

    with st.form("search_form"):
        col1, col2, col3, col4 = st.columns([2, 2, 1, 1])
        with col1:
            keywords = st.text_input("Keywords", value="Business Analyst")
        with col2:
            location = st.text_input("Location", value="Burlington, Ontario")
        with col3:
            country = st.selectbox("Country", ["ca", "us"], format_func=lambda c: "Canada" if c == "ca" else "United States")
        with col4:
            sources = st.multiselect("Sources", ["Adzuna", "Jooble"], default=["Adzuna", "Jooble"])
        searched = st.form_submit_button("Search", type="primary")

    if searched:
        total = 0
        if "Adzuna" in sources:
            with st.spinner("Searching Adzuna..."):
                results = search_adzuna(keywords, location, country=country)
                for job in results:
                    upsert_job_feed_row(conn, {
                        "title": job.get("title", ""),
                        "company": (job.get("company") or {}).get("display_name", "Unknown"),
                        "location": (job.get("location") or {}).get("display_name", ""),
                        "salary": f"{job.get('salary_min', '')}-{job.get('salary_max', '')}" if job.get("salary_min") else "",
                        "url": job.get("redirect_url", ""),
                        "description": job.get("description", ""),
                        "source": "Adzuna",
                    })
                total += len(results)
        if "Jooble" in sources:
            with st.spinner("Searching Jooble..."):
                results = search_jooble(keywords, location)
                for job in results:
                    upsert_job_feed_row(conn, {
                        "title": job.get("title", ""),
                        "company": job.get("company", "Unknown"),
                        "location": job.get("location", ""),
                        "salary": job.get("salary", ""),
                        "url": job.get("link", ""),
                        "description": job.get("snippet", ""),
                        "source": "Jooble",
                    })
                total += len(results)
        st.success(f"Found {total} listings across selected sources. Scroll down to review.")

    with st.expander("Company Watchlist (Greenhouse / Lever)"):
        st.caption(
            "Works only for companies that actually run their careers page on Greenhouse or Lever "
            "(common for tech startups/scaleups). Big banks, Big 4, staffing firms, and government "
            "jobs typically use Workday or proprietary systems instead -- those are in Quick Links below."
        )
        wcol1, wcol2, wcol3, wcol4 = st.columns([2, 1, 2, 1])
        with wcol1:
            wc_name = st.text_input("Company display name", key="wc_name")
        with wcol2:
            wc_platform = st.selectbox("Platform", ["greenhouse", "lever"], key="wc_platform")
        with wcol3:
            wc_token = st.text_input("Board token / slug (from their careers URL)", key="wc_token")
        with wcol4:
            st.write("")
            if st.button("Add to watchlist"):
                if wc_name and wc_token:
                    add_watchlist_company(conn, wc_name, wc_platform, wc_token)
                    st.rerun()

        watchlist_df = get_watchlist_df(conn)
        if not watchlist_df.empty:
            for _, wrow in watchlist_df.iterrows():
                wc1, wc2, wc3 = st.columns([3, 1, 1])
                wc1.write(f"**{wrow['company_name']}** ({wrow['platform']}: `{wrow['board_token']}`)")
                if wc2.button("Fetch jobs", key=f"fetch_wl_{wrow['id']}"):
                    with st.spinner(f"Fetching {wrow['company_name']}..."):
                        if wrow["platform"] == "greenhouse":
                            jobs = fetch_greenhouse_board(wrow["board_token"])
                            for j in jobs:
                                upsert_job_feed_row(conn, {
                                    "title": j.get("title", ""), "company": wrow["company_name"],
                                    "location": (j.get("location") or {}).get("name", ""), "salary": "",
                                    "url": j.get("absolute_url", ""), "description": j.get("content", "") or "",
                                    "source": wrow["company_name"],
                                })
                        else:
                            jobs = fetch_lever_board(wrow["board_token"])
                            for j in jobs:
                                upsert_job_feed_row(conn, {
                                    "title": j.get("text", ""), "company": wrow["company_name"],
                                    "location": (j.get("categories") or {}).get("location", ""), "salary": "",
                                    "url": j.get("hostedUrl", ""), "description": j.get("descriptionPlain", "") or "",
                                    "source": wrow["company_name"],
                                })
                        st.success(f"Fetched {len(jobs)} postings from {wrow['company_name']}.")
                        st.rerun()
                if wc3.button("Remove", key=f"remove_wl_{wrow['id']}"):
                    remove_watchlist_company(conn, wrow["id"])
                    st.rerun()
        else:
            st.caption("No companies watchlisted yet.")

    with st.expander("Quick Links -- LinkedIn, Indeed, Job Bank & named companies"):
        st.caption(
            "These platforms don't offer a public search API, so results aren't pulled into the app -- "
            "these links jump straight to a prefilled search (where the site supports it) so there's "
            "nothing to type."
        )
        import urllib.parse
        kw_enc = urllib.parse.quote(keywords if "keywords" in dir() else "Business Analyst")
        loc_enc = urllib.parse.quote(location if "location" in dir() else "Burlington, Ontario")
        qcols = st.columns(3)
        for i, (label, key, template) in enumerate(QUICK_LINKS):
            url = template.format(kw=kw_enc, loc=loc_enc) if "{kw}" in template or "{loc}" in template else template
            with qcols[i % 3]:
                st.link_button(label, url, use_container_width=True)

    feed_df = get_job_feed_df(conn)
    feed_df = feed_df[feed_df["status"] != "applied"]

    if feed_df.empty:
        st.caption("No listings yet — run a search above.")
    else:
        for _, row in feed_df.iterrows():
            with st.expander(f"{row['title']} — {row['company']} ({row['location']})  [{row['source']} · {row['status']}]"):
                if row["salary"] and row["salary"] != "-":
                    st.write(f"**Salary:** {row['salary']}")
                st.write(row["description"][:1200] + ("..." if len(row["description"]) > 1200 else ""))
                st.link_button("Open full posting ↗", row["url"])

                fit_json = json.loads(row["fit_json"]) if row["fit_json"] else None
                tailor_json = json.loads(row["tailor_json"]) if row["tailor_json"] else None
                cover_json = json.loads(row["cover_json"]) if row["cover_json"] else None

                bcol1, bcol2, bcol3 = st.columns(3)
                with bcol1:
                    if st.button("Analyze Fit", key=f"fit_{row['id']}"):
                        with st.spinner("Scoring fit..."):
                            result = call_tool(
                                "You are an expert technical recruiter. Be honest and specific, not flattering.",
                                f"RESUME:\n{st.session_state.resume_text}\n\nJOB DESCRIPTION:\n{row['description']}\n\nAnalyze fit.",
                                FIT_SCHEMA,
                            )
                            if result:
                                update_job_feed_row(conn, row["id"], fit_json=json.dumps(result),
                                                     fit_score=result.get("fit_score", 0), status="analyzed")
                                st.rerun()
                with bcol2:
                    if st.button("Generate Tailored Resume + Cover Letter", key=f"tailor_{row['id']}"):
                        with st.spinner("Tailoring..."):
                            t = call_tool(
                                "You are an expert resume writer. Rewrite content to match the job description's language. Never invent employers, dates, titles, metrics, or skills not already present in the original resume.",
                                f"ORIGINAL RESUME:\n{st.session_state.resume_text}\n\nJOB DESCRIPTION:\n{row['description']}\n\nProduce tailored headline, summary, and bullet rewrites.",
                                TAILOR_SCHEMA,
                            )
                            c = call_tool(
                                "You write concise, specific, non-generic cover letters (3-4 short paragraphs) grounded only in the candidate's real resume. No invented facts.",
                                f"RESUME:\n{st.session_state.resume_text}\n\nJOB DESCRIPTION:\n{row['description']}\n\nWrite the cover letter.",
                                COVER_LETTER_SCHEMA,
                            )
                            if t and c:
                                update_job_feed_row(conn, row["id"], tailor_json=json.dumps(t),
                                                     cover_json=json.dumps(c), status="tailored")
                                st.rerun()
                with bcol3:
                    if st.button("Mark as Applied", key=f"applied_{row['id']}", disabled=not tailor_json):
                        insert_application(
                            conn,
                            company=row["company"], role=row["title"], url=row["url"],
                            fit_score=row["fit_score"] or 0, status="Applied",
                            date_applied=dt.date.today().isoformat(), date_updated=dt.datetime.now().isoformat(),
                            notes="", jd_text=row["description"],
                            tailor_json=row["tailor_json"] or "", cover_json=row["cover_json"] or "",
                        )
                        update_job_feed_row(conn, row["id"], status="applied")
                        st.success("Logged to Dashboard.")
                        st.rerun()

                if fit_json:
                    score = fit_json.get("fit_score", 0)
                    color = "🟢" if score >= 75 else ("🟡" if score >= 50 else "🔴")
                    st.write(f"{color} **Fit score: {score}/100** — {fit_json.get('recommendation','')}")
                    st.write(fit_json.get("summary", ""))
                if tailor_json:
                    st.write("**Tailored headline:** " + tailor_json.get("tailored_headline", ""))
                    st.write("**Tailored summary:** " + tailor_json.get("tailored_summary", ""))
                    resume_buf = export_tailored_resume_docx(tailor_json)
                    st.download_button("Download tailored resume (.docx)", data=resume_buf,
                                        file_name=f"Kahini_Resume_{row['company']}.docx", key=f"dl_resume_{row['id']}")
                if cover_json:
                    cover_buf = export_cover_letter_docx(cover_json)
                    st.download_button("Download cover letter (.docx)", data=cover_buf,
                                        file_name=f"Kahini_CoverLetter_{row['company']}.docx", key=f"dl_cover_{row['id']}")

    st.caption(
        "Note: needs `ADZUNA_APP_ID` and `ADZUNA_APP_KEY` in Secrets — free account at developer.adzuna.com. "
        "Applying still happens on the real site — 'Open full posting' takes you there; 'Mark as Applied' just logs it here."
    )

# --- Page 3: Analyze Fit (manual paste) ---
elif page == "3. Analyze Fit":
    st.subheader("Paste a Job Posting")
    input_mode = st.radio("Input method", ["Paste job description text", "Fetch from URL"], horizontal=True)

    if input_mode == "Fetch from URL":
        url = st.text_input("Job posting URL")
        if st.button("Fetch posting"):
            fetched = fetch_job_posting(url)
            if fetched:
                st.session_state.jd_text = fetched
                st.success("Fetched. Review below — some job boards render JDs via JavaScript and won't extract cleanly.")

    st.session_state.jd_text = st.text_area("Job description", value=st.session_state.jd_text, height=300)

    if st.button("Analyze Fit", type="primary"):
        if not st.session_state.jd_text.strip():
            st.warning("Paste or fetch a job description first.")
        else:
            with st.spinner("Analyzing fit against resume..."):
                st.session_state.fit_result = call_tool(
                    "You are an expert technical recruiter. Be honest and specific, not flattering.",
                    f"RESUME:\n{st.session_state.resume_text}\n\nJOB DESCRIPTION:\n{st.session_state.jd_text}\n\nAnalyze fit.",
                    FIT_SCHEMA,
                )

    if st.session_state.fit_result:
        r = st.session_state.fit_result
        score = r.get("fit_score", 0)
        color = "🟢" if score >= 75 else ("🟡" if score >= 50 else "🔴")
        st.metric("Fit Score", f"{score}/100")
        st.write(f"{color} **{r.get('recommendation','')}**")
        st.write(f"**Role:** {r.get('role_title','')}  |  **Company:** {r.get('company_name','')}")
        st.write(r.get("summary", ""))
        col1, col2 = st.columns(2)
        with col1:
            st.write("**Matched skills**")
            for s in r.get("matched_skills", []):
                st.write(f"- {s}")
        with col2:
            st.write("**Gaps**")
            for s in r.get("missing_skills", []):
                st.write(f"- {s}")

# --- Page 4: Tailor Resume ---
elif page == "4. Tailor Resume":
    st.subheader("Tailor Resume to This Job")
    if not st.session_state.jd_text.strip():
        st.warning("Add a job description in the Analyze Fit tab first.")
    else:
        if st.button("Generate Tailored Resume", type="primary"):
            with st.spinner("Tailoring resume..."):
                st.session_state.tailor_result = call_tool(
                    "You are an expert resume writer. Rewrite content to match the job description's language. Never invent employers, dates, titles, metrics, or skills not already present in the original resume.",
                    f"ORIGINAL RESUME:\n{st.session_state.resume_text}\n\nJOB DESCRIPTION:\n{st.session_state.jd_text}\n\nProduce tailored headline, summary, and bullet rewrites.",
                    TAILOR_SCHEMA,
                )

        if st.session_state.tailor_result:
            t = st.session_state.tailor_result
            st.write("**Headline:** " + t.get("tailored_headline", ""))
            st.write("**Summary:** " + t.get("tailored_summary", ""))
            for entry in t.get("experience_rewrites", []):
                st.write(f"**{entry.get('role_company','')}**")
                for b in entry.get("bullets", []):
                    st.write(f"- {b}")

            buf = export_tailored_resume_docx(t)
            st.download_button("Download tailored resume (.docx)", data=buf,
                                file_name=f"Kahini_Gandhi_Resume_{dt.date.today().isoformat()}.docx")

# --- Page 5: Cover Letter ---
elif page == "5. Cover Letter":
    st.subheader("Cover Letter")
    if not st.session_state.jd_text.strip():
        st.warning("Add a job description in the Analyze Fit tab first.")
    else:
        if st.button("Generate Cover Letter", type="primary"):
            with st.spinner("Drafting cover letter..."):
                st.session_state.cover_result = call_tool(
                    "You write concise, specific, non-generic cover letters (3-4 short paragraphs) grounded only in the candidate's real resume. No invented facts.",
                    f"RESUME:\n{st.session_state.resume_text}\n\nJOB DESCRIPTION:\n{st.session_state.jd_text}\n\nWrite the cover letter.",
                    COVER_LETTER_SCHEMA,
                )

        if st.session_state.cover_result:
            c = st.session_state.cover_result
            st.write(c.get("greeting", ""))
            for para in c.get("body_paragraphs", []):
                st.write(para)
            st.write(c.get("closing", ""))

            buf = export_cover_letter_docx(c)
            st.download_button("Download cover letter (.docx)", data=buf,
                                file_name=f"Kahini_Gandhi_CoverLetter_{dt.date.today().isoformat()}.docx")

        if st.session_state.fit_result and st.session_state.tailor_result and st.session_state.cover_result:
            st.divider()
            if st.button("Save this application to Dashboard"):
                r = st.session_state.fit_result
                insert_application(
                    conn,
                    company=r.get("company_name", "Unknown"), role=r.get("role_title", "Unknown"),
                    url="", fit_score=r.get("fit_score", 0), status="Applied",
                    date_applied=dt.date.today().isoformat(), date_updated=dt.datetime.now().isoformat(),
                    notes="", jd_text=st.session_state.jd_text,
                    tailor_json=json.dumps(st.session_state.tailor_result),
                    cover_json=json.dumps(st.session_state.cover_result),
                )
                st.success("Saved to Dashboard.")

# --- Page 6: Dashboard ---
elif page == "6. Dashboard":
    st.subheader("Application Dashboard")

    apps_df = get_applications_df(conn)

    if apps_df.empty:
        st.caption("No applications logged yet. Use Job Feed or Cover Letter tab to log one.")
    else:
        m1, m2, m3, m4, m5 = st.columns(5)
        m1.metric("Total", len(apps_df))
        m2.metric("Applied", (apps_df["status"] == "Applied").sum())
        m3.metric("Interviewing", (apps_df["status"] == "Interviewing").sum())
        m4.metric("Offers", (apps_df["status"] == "Offer").sum())
        m5.metric("Rejected", (apps_df["status"] == "Rejected").sum())

        st.bar_chart(apps_df["status"].value_counts())

        st.divider()
        st.write("**All applications** (edit status/notes directly in the table)")

        editable = apps_df[["id", "company", "role", "fit_score", "status", "date_applied", "notes"]].copy()
        edited = st.data_editor(
            editable,
            column_config={
                "id": st.column_config.NumberColumn("ID", disabled=True),
                "status": st.column_config.SelectboxColumn(
                    "Status", options=["Not applied", "Applied", "Interviewing", "Offer", "Rejected", "Withdrawn"]
                ),
            },
            use_container_width=True,
            hide_index=True,
            key="dashboard_editor",
        )

        if st.button("Save changes"):
            for _, row in edited.iterrows():
                original = editable[editable["id"] == row["id"]].iloc[0]
                if row["status"] != original["status"] or row["notes"] != original["notes"]:
                    update_application(conn, row["id"], status=row["status"], notes=row["notes"],
                                        date_updated=dt.datetime.now().isoformat())
            st.success("Saved.")
            st.rerun()

        st.divider()
        st.write("**View the tailored resume / cover letter used for a specific application**")
        for _, row in apps_df.iterrows():
            with st.expander(f"{row['company']} — {row['role']}  (applied {row['date_applied']}, {row['status']})"):
                if row["tailor_json"]:
                    t = json.loads(row["tailor_json"])
                    st.write("**Tailored headline:** " + t.get("tailored_headline", ""))
                    st.write("**Tailored summary:** " + t.get("tailored_summary", ""))
                    for entry in t.get("experience_rewrites", []):
                        st.write(f"*{entry.get('role_company','')}*")
                        for b in entry.get("bullets", []):
                            st.write(f"- {b}")
                    resume_buf = export_tailored_resume_docx(t)
                    st.download_button("Download this tailored resume (.docx)", data=resume_buf,
                                        file_name=f"Kahini_Resume_{row['company']}.docx", key=f"dash_resume_{row['id']}")
                if row["cover_json"]:
                    c = json.loads(row["cover_json"])
                    cover_buf = export_cover_letter_docx(c)
                    st.download_button("Download this cover letter (.docx)", data=cover_buf,
                                        file_name=f"Kahini_CoverLetter_{row['company']}.docx", key=f"dash_cover_{row['id']}")
                if row["url"]:
                    st.link_button("Original posting ↗", row["url"])

        buf = io.BytesIO()
        apps_df.drop(columns=["tailor_json", "cover_json", "jd_text"]).to_excel(buf, index=False, engine="openpyxl")
        buf.seek(0)
        st.download_button("Download full tracker (.xlsx)", data=buf, file_name="Kahini_Application_Tracker.xlsx")

# --- Page 7: Email Check (placeholder) ---
elif page == "7. Email Check":
    st.subheader("Email Check — Coming Soon")
    st.info(
        "This is on the backlog: an optional, read-only scan of Kahini's Gmail inbox (via a Gmail "
        "'app password') to flag likely responses to logged applications — interview invites, "
        "rejections, or requests for more info — so she can confirm and update status on the "
        "Dashboard. Not built yet; status updates are manual for now via the Dashboard tab."
    )
